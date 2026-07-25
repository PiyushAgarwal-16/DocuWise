"""
core/extractor.py — Text extraction layer for DocuWise.

Responsibilities:
  - Extract plain text from PDF, DOCX, TXT, and image files.
  - HYBRID PDF extraction (Phase 4): use native text where present, OCR only the
    pages that lack text, then merge — never OCR a page that already has text.
  - Image files are first-class documents (Phase 5): extracted entirely via OCR.
  - Compute MD5 of the raw file bytes for content-addressed caching.
  - Consult the Content Cache (Phase 2) right after MD5:
        * complete hit  → restore analysis + embedding, skip OCR/NVIDIA/embedding
        * text hit      → reuse cached final/OCR text, skip OCR only
        * miss          → extract (+ OCR), persist, seed the cache
  - Persist extraction results and extraction metadata (Phase 6) to the database.
  - Never let OCR failure crash a scan (Phase 12).

This module does NOT perform:
  - NVIDIA analysis (see analyzer.py)
  - Embedding generation (see embedder.py)
  - PPTX / XLSX extraction (out of scope)
  - Any UI interaction
"""

import hashlib
import logging
import re
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from config import (
    CONTENT_CACHE_ENABLED,
    IMAGE_EXTENSIONS,
    MIN_WORD_COUNT,
    OCR_ENABLED,
    OCR_ENGINE,
    OCR_MAX_PDF_PAGES,
    OCR_MIN_TEXT_THRESHOLD,
    OCR_RENDER_DPI,
    OCR_TIMEOUT_SECONDS,
    OCR_VERSION,
)
from core import content_cache
from core.database import update_document_extraction, update_document_status
from core.logging_events import log_event
from core.ocr import BaseOCREngine, get_ocr_engine

logger = logging.getLogger(__name__)

_ = MIN_WORD_COUNT  # retained import: threshold used by analyzer/UI, kept for parity


# ---------------------------------------------------------------------------
# Result dataclass
# ---------------------------------------------------------------------------

@dataclass
class ExtractedDocument:
    """
    Structured result returned by extract_document() and process_document().

    Attributes:
        file_path:     Absolute path to the source file.
        text:          Final plain-text content used downstream (native+OCR merge).
        word_count:    Whitespace-delimited token count of *text*.
        page_count:    Number of pages (PDF/image). None for DOCX and TXT.
        success:       True if extraction completed without a fatal error.
        image_only:    True when the document has no usable text (OCR disabled,
                       failed, or produced nothing) — skips analysis/embedding.
        cache_hit:     True when a complete Content Cache entry was restored,
                       meaning the pipeline should skip analysis AND embedding.
        error_message: Human-readable failure description, or None.

        md5_hash:            MD5 of the raw file bytes (set once computed).
        extraction_method:   native | hybrid | ocr_only | image_only (Phase 6).
        ocr_engine/version:  Provenance of any OCR performed.
        ocr_confidence:      Mean OCR confidence 0.0–1.0, or None.
        ocr_pages_processed: Pages actually OCR'd.
        ocr_pages_skipped:   Pages that already had native text (not OCR'd).
        ocr_processing_time_ms: Total OCR wall-clock time in milliseconds.
        ocr_cached:          1 if OCR/text was reused from the cache.
    """
    file_path: str
    text: str = ""
    word_count: int = 0
    page_count: Optional[int] = None
    success: bool = False
    image_only: bool = False
    cache_hit: bool = False
    error_message: Optional[str] = None

    # Extraction metadata (Phase 6)
    md5_hash: Optional[str] = None
    extraction_method: Optional[str] = None
    ocr_engine: Optional[str] = None
    ocr_version: Optional[str] = None
    ocr_confidence: Optional[float] = None
    ocr_pages_processed: int = 0
    ocr_pages_skipped: int = 0
    ocr_processing_time_ms: int = 0
    ocr_cached: int = 0


@dataclass
class _OCROutcome:
    """Internal carrier for the result of native+OCR extraction of one file."""
    final_text: str = ""
    native_text: str = ""
    ocr_text: Optional[str] = None
    extraction_method: str = "native"
    ocr_confidence: Optional[float] = None
    ocr_pages_processed: int = 0
    ocr_pages_skipped: int = 0
    ocr_processing_time_ms: int = 0
    page_count: Optional[int] = None
    success: bool = True
    error_message: Optional[str] = None


# ---------------------------------------------------------------------------
# MD5 helper
# ---------------------------------------------------------------------------

def compute_md5(file_path: str) -> str:
    """
    Compute the MD5 hex-digest of the raw bytes of a file.

    Reads in 64 KB chunks so large files never load fully into memory. This hash
    is the primary key of the Content Cache — identical content (moved / copied /
    duplicate files) collides here and reuses all cached work.

    Args:
        file_path: Absolute path to the file.

    Returns:
        Lowercase 32-character hex-digest string.

    Raises:
        OSError: If the file cannot be opened or read.
    """
    hasher = hashlib.md5()
    chunk_size = 65_536  # 64 KB

    with open(file_path, "rb") as fh:
        while chunk := fh.read(chunk_size):
            hasher.update(chunk)

    return hasher.hexdigest()


# ---------------------------------------------------------------------------
# Word count helper
# ---------------------------------------------------------------------------

def _count_words(text: str) -> int:
    """Count whitespace-delimited words in *text*; 0 for empty/whitespace."""
    stripped = text.strip()
    if not stripped:
        return 0
    return len(re.split(r"\s+", stripped))


# ---------------------------------------------------------------------------
# Native PDF extraction
# ---------------------------------------------------------------------------

def _pdf_native_pages(file_path: str) -> tuple[list[str], Optional[str]]:
    """
    Read native (embedded) text from every page of a PDF.

    Args:
        file_path: Absolute path to the PDF.

    Returns:
        (page_texts, error_message). page_texts has one entry per page (possibly
        empty strings for image-only pages). error_message is None on success or
        a description if the PDF is corrupt / encrypted-and-locked.
    """
    try:
        doc = fitz.open(file_path)
    except fitz.FileDataError as exc:
        return [], f"Corrupt or unreadable PDF: {exc}"
    except Exception as exc:  # noqa: BLE001
        return [], f"Unexpected PDF open error: {exc}"

    try:
        if doc.is_encrypted and not doc.authenticate(""):
            return [], "PDF is encrypted and password-protected."

        page_texts = [(page.get_text("text") or "").strip() for page in doc]
        return page_texts, None
    except Exception as exc:  # noqa: BLE001
        return [], f"Unexpected PDF extraction error: {exc}"
    finally:
        doc.close()


def extract_pdf(file_path: str) -> ExtractedDocument:
    """
    Extract native (non-OCR) plain text from a PDF using PyMuPDF.

    Retained as the public native extractor for backward compatibility. The
    hybrid OCR path is applied later in process_document(); this function only
    reports the text embedded in the PDF.

    Returns:
        ExtractedDocument (extraction_method='native') on success, or
        success=False with error_message on failure.
    """
    page_texts, error = _pdf_native_pages(file_path)
    if error is not None:
        logger.warning("PDF extraction failed for '%s': %s", file_path, error)
        return ExtractedDocument(file_path=file_path, success=False, error_message=error)

    full_text = "\n\n".join(p for p in page_texts if p)
    return ExtractedDocument(
        file_path=file_path,
        text=full_text,
        word_count=_count_words(full_text),
        page_count=len(page_texts),
        success=True,
        extraction_method="native",
    )


def extract_docx(file_path: str) -> ExtractedDocument:
    """
    Extract plain text from a DOCX file using python-docx.

    Returns:
        ExtractedDocument (extraction_method='native') on success, or
        success=False with error_message on failure. page_count is always None.
    """
    try:
        docx = DocxDocument(file_path)
        paragraphs = [p.text.strip() for p in docx.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)
        return ExtractedDocument(
            file_path=file_path,
            text=full_text,
            word_count=_count_words(full_text),
            page_count=None,
            success=True,
            extraction_method="native",
        )
    except Exception as exc:  # noqa: BLE001
        msg = f"DOCX extraction error: {exc}"
        logger.error("DOCX extraction failed for '%s': %s", file_path, exc, exc_info=True)
        return ExtractedDocument(file_path=file_path, success=False, error_message=msg)


def extract_txt(file_path: str) -> ExtractedDocument:
    """
    Extract plain text from a TXT file, trying UTF-8 then latin-1.

    Returns:
        ExtractedDocument (extraction_method='native') on success, or
        success=False with error_message on failure. page_count is always None.
    """
    for encoding in ("utf-8", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding, errors="replace") as fh:
                full_text = fh.read()
            return ExtractedDocument(
                file_path=file_path,
                text=full_text,
                word_count=_count_words(full_text),
                page_count=None,
                success=True,
                extraction_method="native",
            )
        except UnicodeDecodeError:
            continue
        except OSError as exc:
            msg = f"Could not read TXT file: {exc}"
            logger.error("TXT extraction failed for '%s': %s", file_path, exc)
            return ExtractedDocument(file_path=file_path, success=False, error_message=msg)

    msg = "All encoding attempts failed."
    logger.error("TXT extraction failed for '%s': %s", file_path, msg)
    return ExtractedDocument(file_path=file_path, success=False, error_message=msg)


# ---------------------------------------------------------------------------
# Native dispatcher (backward-compatible public API)
# ---------------------------------------------------------------------------

_EXTRACTOR_MAP = {
    ".pdf":  extract_pdf,
    ".docx": extract_docx,
    ".txt":  extract_txt,
}


def extract_document(file_path: str) -> ExtractedDocument:
    """
    Dispatch to the correct NATIVE extractor based on file extension.

    Image files return an empty-text success result (extraction is entirely OCR,
    performed later). Unsupported extensions return success=False.

    Args:
        file_path: Absolute path to the document.

    Returns:
        ExtractedDocument carrying native text only (no OCR applied here).
    """
    path = Path(file_path)

    if not path.exists():
        return ExtractedDocument(file_path=file_path, success=False,
                                 error_message=f"File not found: '{file_path}'")
    if not path.is_file():
        return ExtractedDocument(file_path=file_path, success=False,
                                 error_message=f"Path is not a regular file: '{file_path}'")

    extension = path.suffix.lower()

    if extension in IMAGE_EXTENSIONS:
        # Images have no native text — OCR happens in process_document().
        return ExtractedDocument(file_path=file_path, text="", page_count=1, success=True)

    extractor_fn = _EXTRACTOR_MAP.get(extension)
    if extractor_fn is None:
        return ExtractedDocument(file_path=file_path, success=False,
                                 error_message=f"Unsupported file extension: '{extension}'")
    return extractor_fn(file_path)


# ---------------------------------------------------------------------------
# OCR helpers (Phase 4/5) — only invoked on a Content Cache miss
# ---------------------------------------------------------------------------

def _ocr_enabled_engine() -> Optional[BaseOCREngine]:
    """Return a ready OCR engine, or None if OCR is disabled/unavailable."""
    if not OCR_ENABLED:
        return None
    engine = get_ocr_engine()
    return engine if engine.is_available() else None


def _render_pdf_page(page, dpi: int):
    """Render a PyMuPDF page to a PIL RGB image at *dpi* (no alpha channel)."""
    from PIL import Image

    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=matrix, alpha=False)
    try:
        return Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    finally:
        pix = None  # release the pixmap buffer promptly (Phase 8 — memory)


def _extract_pdf_with_ocr(file_path: str, page_texts: list[str], filename: str) -> _OCROutcome:
    """
    Hybrid PDF extraction: keep native text, OCR only text-poor pages, merge.

    Args:
        file_path:  Absolute path to the PDF.
        page_texts: Native text per page (from _pdf_native_pages).
        filename:   Bare filename, for structured logs.

    Returns:
        _OCROutcome with merged final_text and full extraction metadata. Never
        raises — OCR errors are contained per page.
    """
    outcome = _OCROutcome(page_count=len(page_texts))
    engine = _ocr_enabled_engine()

    # Which pages need OCR? Those below the native-text threshold.
    poor_pages = [i for i, t in enumerate(page_texts) if len(t) < OCR_MIN_TEXT_THRESHOLD]
    native_present = any(len(t) >= OCR_MIN_TEXT_THRESHOLD for t in page_texts)

    # No OCR possible/needed → native-only result.
    if engine is None or not poor_pages:
        outcome.native_text = "\n\n".join(t for t in page_texts if t)
        outcome.final_text = outcome.native_text
        outcome.ocr_pages_skipped = sum(1 for t in page_texts if t)
        outcome.extraction_method = "native" if outcome.final_text.strip() else "image_only"
        if poor_pages and engine is None:
            log_event(logger, "OCR_PAGE_SKIPPED", file=filename,
                      pages=len(poor_pages), reason="ocr_unavailable")
        return outcome

    log_event(logger, "OCR_STARTED", file=filename, engine=engine.name,
              pages=len(poor_pages), total_pages=len(page_texts))

    final_parts: list[str] = []
    native_parts: list[str] = []
    ocr_parts: list[str] = []
    confidences: list[float] = []
    start = time.time()

    try:
        doc = fitz.open(file_path)
    except Exception as exc:  # noqa: BLE001
        # Native text was already read successfully; degrade to native-only.
        logger.warning("OCR re-open failed for '%s': %s — using native text.", filename, exc)
        outcome.native_text = "\n\n".join(t for t in page_texts if t)
        outcome.final_text = outcome.native_text
        outcome.extraction_method = "native" if outcome.final_text.strip() else "image_only"
        return outcome

    try:
        if doc.is_encrypted:
            doc.authenticate("")

        for i in range(len(page_texts)):
            native = page_texts[i]

            if len(native) >= OCR_MIN_TEXT_THRESHOLD:
                # Page already has text — never OCR it (Phase 8).
                final_parts.append(native)
                native_parts.append(native)
                outcome.ocr_pages_skipped += 1
                continue

            # Budget / cap guards.
            over_budget = (time.time() - start) >= OCR_TIMEOUT_SECONDS
            over_cap = outcome.ocr_pages_processed >= OCR_MAX_PDF_PAGES
            if over_budget or over_cap:
                if native:
                    final_parts.append(native)
                outcome.ocr_pages_skipped += 1
                log_event(logger, "OCR_PAGE_SKIPPED", file=filename, page=i + 1,
                          reason="timeout" if over_budget else "page_cap")
                continue

            try:
                image = _render_pdf_page(doc[i], OCR_RENDER_DPI)
                result = engine.recognize(image)
            except Exception as exc:  # noqa: BLE001
                result = None
                log_event(logger, "OCR_FAILED", level=logging.WARNING,
                          file=filename, page=i + 1, error=str(exc))

            if result is not None and result.success and result.text.strip():
                ocr_parts.append(result.text)
                final_parts.append(result.text)
                confidences.append(result.confidence)
                outcome.ocr_pages_processed += 1
                log_event(logger, "OCR_PAGE_COMPLETED", file=filename, page=i + 1,
                          confidence=round(result.confidence, 3),
                          lines=result.line_count)
            else:
                # OCR yielded nothing usable — keep whatever native text existed.
                if native:
                    final_parts.append(native)
                outcome.ocr_pages_skipped += 1
                if result is not None and not result.success:
                    log_event(logger, "OCR_FAILED", level=logging.WARNING,
                              file=filename, page=i + 1, error=result.error_message)
    finally:
        doc.close()

    outcome.ocr_processing_time_ms = int((time.time() - start) * 1000)
    outcome.native_text = "\n\n".join(native_parts)
    outcome.ocr_text = "\n\n".join(ocr_parts) if ocr_parts else None
    outcome.final_text = "\n\n".join(final_parts)
    outcome.ocr_confidence = (sum(confidences) / len(confidences)) if confidences else None

    # Classify extraction method.
    if outcome.ocr_pages_processed > 0 and native_present:
        outcome.extraction_method = "hybrid"
    elif outcome.ocr_pages_processed > 0:
        outcome.extraction_method = "ocr_only"
    elif outcome.final_text.strip():
        outcome.extraction_method = "native"
    else:
        outcome.extraction_method = "image_only"

    log_event(logger, "OCR_COMPLETED", file=filename,
              pages_ocred=outcome.ocr_pages_processed,
              pages_skipped=outcome.ocr_pages_skipped,
              duration_ms=outcome.ocr_processing_time_ms,
              confidence=(round(outcome.ocr_confidence, 3)
                          if outcome.ocr_confidence is not None else None),
              method=outcome.extraction_method)
    return outcome


def _extract_image_with_ocr(file_path: str, filename: str) -> _OCROutcome:
    """
    OCR an image file (Phase 5). Images are first-class, text-free documents.

    Returns:
        _OCROutcome with method 'ocr_only' on success or 'image_only' if OCR is
        unavailable / fails / finds no text. Never raises.
    """
    outcome = _OCROutcome(page_count=1, native_text="")
    engine = _ocr_enabled_engine()

    if engine is None:
        outcome.extraction_method = "image_only"
        outcome.final_text = ""
        log_event(logger, "OCR_PAGE_SKIPPED", file=filename, reason="ocr_unavailable")
        return outcome

    log_event(logger, "OCR_STARTED", file=filename, engine=engine.name, pages=1)
    start = time.time()

    try:
        from PIL import Image
        with Image.open(file_path) as im:
            im.load()
            rgb = im.convert("RGB")
        result = engine.recognize(rgb)
    except Exception as exc:  # noqa: BLE001
        outcome.ocr_processing_time_ms = int((time.time() - start) * 1000)
        outcome.extraction_method = "image_only"
        log_event(logger, "OCR_FAILED", level=logging.WARNING, file=filename, error=str(exc))
        return outcome

    outcome.ocr_processing_time_ms = int((time.time() - start) * 1000)

    if result.success and result.text.strip():
        outcome.final_text = result.text
        outcome.ocr_text = result.text
        outcome.ocr_pages_processed = 1
        outcome.ocr_confidence = result.confidence
        outcome.extraction_method = "ocr_only"
        log_event(logger, "OCR_COMPLETED", file=filename, pages_ocred=1,
                  duration_ms=outcome.ocr_processing_time_ms,
                  confidence=round(result.confidence, 3),
                  method="ocr_only")
    else:
        outcome.extraction_method = "image_only"
        log_event(logger, "OCR_FAILED", level=logging.WARNING, file=filename,
                  error=(result.error_message or "no text recognised"))

    return outcome


# ---------------------------------------------------------------------------
# Pipeline integration — full single-document extraction
# ---------------------------------------------------------------------------

def process_document(file_path: str) -> ExtractedDocument:
    """
    Full extraction workflow for a single document (Phase 4/5/9).

    Flow:
        1. Native extraction (PDF/DOCX/TXT native text; images → empty).
        2. Compute MD5 of the raw bytes.
        3. Content Cache lookup:
             - complete hit → restore analysis + embedding, mark cache_hit,
               return (pipeline skips analysis AND embedding).
             - text hit     → reuse cached final/OCR text, skip OCR only.
             - miss         → run OCR where required (hybrid PDF / image).
        4. If no usable text → mark 'image_only' (not a failure).
        5. Persist final text + extraction metadata; seed the Content Cache.

    Safe to call repeatedly — all database writes are keyed by file_path.

    Returns:
        ExtractedDocument. Inspect .success, .cache_hit, and .image_only to
        decide how to proceed.
    """
    logger.info("Processing: '%s'", file_path)
    path = Path(file_path)
    filename = path.name
    extension = path.suffix.lower()
    is_image = extension in IMAGE_EXTENSIONS

    # ── Step 1: Native extraction ────────────────────────────────────────────
    native = extract_document(file_path)
    if not native.success:
        logger.warning("Extraction failed for '%s': %s", filename, native.error_message)
        update_document_status(file_path, "failed", error_message=native.error_message)
        return native

    # For PDFs we need per-page native text to drive hybrid OCR.
    pdf_page_texts: Optional[list[str]] = None
    if extension == ".pdf":
        pdf_page_texts, pdf_error = _pdf_native_pages(file_path)
        if pdf_error is not None:
            update_document_status(file_path, "failed", error_message=pdf_error)
            return ExtractedDocument(file_path=file_path, success=False, error_message=pdf_error)

    # ── Step 2: MD5 ──────────────────────────────────────────────────────────
    try:
        md5_hash = compute_md5(file_path)
    except OSError as exc:
        msg = f"MD5 computation failed: {exc}"
        logger.error("Hashing failed for '%s': %s", file_path, exc)
        update_document_status(file_path, "failed", error_message=msg)
        return ExtractedDocument(file_path=file_path, success=False, error_message=msg)
    native.md5_hash = md5_hash

    # ── Step 3: Content Cache lookup ─────────────────────────────────────────
    reuse: Optional[dict] = None
    if CONTENT_CACHE_ENABLED:
        entry = content_cache.lookup(md5_hash, exclude_file_path=file_path, filename=filename)
        if entry is not None and content_cache.is_complete(entry):
            # Full hit — restore everything, skip OCR + NVIDIA + embedding.
            content_cache.restore_to_document(file_path, entry)
            native.cache_hit = True
            native.image_only = False
            native.text = entry.get("final_text") or ""
            native.word_count = _count_words(native.text)
            native.extraction_method = entry.get("extraction_method") or native.extraction_method
            native.ocr_engine = entry.get("ocr_engine")
            native.ocr_version = entry.get("ocr_version")
            native.ocr_confidence = entry.get("ocr_confidence")
            native.ocr_cached = 1
            return native
        if entry is not None and content_cache.has_text(entry):
            reuse = entry  # text hit — reuse cached text, skip OCR only

    # ── Step 4: Build final text (reuse cache, or extract + OCR) ─────────────
    if reuse is not None:
        outcome = _OCROutcome(
            final_text=reuse.get("final_text") or "",
            native_text=reuse.get("native_text") or "",
            ocr_text=reuse.get("ocr_text"),
            extraction_method=reuse.get("extraction_method")
                              or ("ocr_only" if is_image else "native"),
            ocr_confidence=reuse.get("ocr_confidence"),
            ocr_pages_processed=reuse.get("ocr_pages_processed") or 0,
            page_count=1 if is_image else (len(pdf_page_texts) if pdf_page_texts else None),
        )
        # Only an OCR-cache reuse if the cached content actually involved OCR.
        ocr_cached = 1 if (reuse.get("ocr_text") or reuse.get("ocr_pages_processed")) else 0
    else:
        if is_image:
            outcome = _extract_image_with_ocr(file_path, filename)
        elif extension == ".pdf":
            outcome = _extract_pdf_with_ocr(file_path, pdf_page_texts or [], filename)
        else:
            # DOCX / TXT — native only, no OCR.
            outcome = _OCROutcome(
                final_text=native.text,
                native_text=native.text,
                extraction_method="native" if native.text.strip() else "image_only",
                page_count=native.page_count,
            )
        ocr_cached = 0

    final_text = outcome.final_text or ""
    engine_name = OCR_ENGINE if outcome.ocr_pages_processed or ocr_cached else None
    engine_version = OCR_VERSION if outcome.ocr_pages_processed or ocr_cached else None

    # ── Step 4b: No usable text → image_only (not a failure, Phase 12) ───────
    if not final_text.strip():
        image_only_msg = "No extractable text (OCR disabled, failed, or empty)."
        log_event(logger, "OCR_FAILED", level=logging.INFO, file=filename,
                  md5=md5_hash, reason="no_text", method="image_only")
        try:
            update_document_extraction(
                file_path=file_path, word_count=0, md5_hash=md5_hash, extracted_text="",
                extraction_method="image_only",
                ocr_engine=engine_name, ocr_version=engine_version,
                ocr_pages_processed=outcome.ocr_pages_processed,
                ocr_pages_skipped=outcome.ocr_pages_skipped,
                ocr_processing_time_ms=outcome.ocr_processing_time_ms,
            )
            update_document_status(file_path, "image_only", error_message=image_only_msg)
        except Exception as exc:  # noqa: BLE001
            logger.error("DB write failed for image-only '%s': %s", file_path, exc)

        native.image_only = True
        native.success = True
        native.text = ""
        native.extraction_method = "image_only"
        native.ocr_pages_skipped = outcome.ocr_pages_skipped
        native.ocr_processing_time_ms = outcome.ocr_processing_time_ms
        return native

    # ── Step 5: Persist success + seed the Content Cache ─────────────────────
    word_count = _count_words(final_text)
    try:
        update_document_extraction(
            file_path=file_path,
            word_count=word_count,
            md5_hash=md5_hash,
            extracted_text=final_text,
            extraction_method=outcome.extraction_method,
            ocr_engine=engine_name,
            ocr_version=engine_version,
            ocr_confidence=outcome.ocr_confidence,
            ocr_pages_processed=outcome.ocr_pages_processed,
            ocr_pages_skipped=outcome.ocr_pages_skipped,
            ocr_processing_time_ms=outcome.ocr_processing_time_ms,
            ocr_cached=ocr_cached,
        )
    except Exception as exc:  # noqa: BLE001
        msg = f"Database write failed after extraction: {exc}"
        logger.error("DB write failed for '%s': %s", file_path, exc, exc_info=True)
        update_document_status(file_path, "failed", error_message=msg)
        return ExtractedDocument(file_path=file_path, success=False, error_message=msg)

    # Seed text layers so identical content never re-OCRs (skips finalise later
    # only if analysis+embedding complete — handled by the pipeline).
    content_cache.seed_text(
        md5_hash,
        native_text=outcome.native_text,
        ocr_text=outcome.ocr_text,
        final_text=final_text,
        extraction_method=outcome.extraction_method,
        ocr_engine=engine_name,
        ocr_version=engine_version,
        ocr_confidence=outcome.ocr_confidence,
        ocr_pages_processed=outcome.ocr_pages_processed,
        ocr_processing_time_ms=outcome.ocr_processing_time_ms,
    )

    logger.info(
        "Extracted '%s' | method=%s | words=%d | md5=%s… | ocr_pages=%d cached=%d",
        filename, outcome.extraction_method, word_count, md5_hash[:8],
        outcome.ocr_pages_processed, ocr_cached,
    )

    native.success = True
    native.text = final_text
    native.word_count = word_count
    native.page_count = outcome.page_count
    native.extraction_method = outcome.extraction_method
    native.ocr_engine = engine_name
    native.ocr_version = engine_version
    native.ocr_confidence = outcome.ocr_confidence
    native.ocr_pages_processed = outcome.ocr_pages_processed
    native.ocr_pages_skipped = outcome.ocr_pages_skipped
    native.ocr_processing_time_ms = outcome.ocr_processing_time_ms
    native.ocr_cached = ocr_cached
    return native
